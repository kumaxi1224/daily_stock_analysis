import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from opencc import OpenCC
import markdown
from bs4 import BeautifulSoup

def process_and_format():
    # 讀取剛才合併好的大報告
    md_file = "reports/final_combined_report.md"
    if not os.path.exists(md_file):
        print("未找到合併後的報告檔案")
        return
        
    with open(md_file, 'r', encoding='utf-8') as f:
        content_sc = f.read()

    # 簡體轉繁體 (使用 opencc 將簡體轉為台灣正體)
    cc = OpenCC('s2twp')
    content_tc = cc.convert(content_sc)

    # 將 Markdown 轉為 HTML 以利解析
    html = markdown.markdown(content_tc, extensions=['tables'])
    soup = BeautifulSoup(html, 'html.parser')

    # 建立 Word 文件並設定「微軟正黑體」與字型大小
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft JhengHei'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')

    for element in soup:
        if element.name in ['h1', 'h2', 'h3', 'h4']:
            level = int(element.name[1]) - 1
            level = max(0, min(level, 8)) # 避免標題層級越界
            heading = doc.add_heading(element.text, level=level)
            # 設定標題也是微軟正黑體與深藍色
            heading.style.font.name = 'Microsoft JhengHei'
            heading.style.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            heading.style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft JhengHei')
            
        elif element.name == 'p':
            doc.add_paragraph(element.text)
            
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Bullet')
                
        elif element.name == 'table':
            rows = element.find_all('tr')
            if rows:
                cols = len(rows[0].find_all(['th', 'td']))
                table = doc.add_table(rows=0, cols=cols)
                table.style = 'Light Shading Accent 1'
                for row in rows:
                    cells = row.find_all(['th', 'td'])
                    row_obj = table.add_row()
                    for i, cell in enumerate(cells):
                        row_obj.cells[i].text = cell.text

    # 儲存為最終的 docx
    out_file = "reports/Daily_Stock_Analysis_Report.docx"
    doc.save(out_file)
    print(f"已成功轉換並排版為繁體 Word 檔: {out_file}")

if __name__ == "__main__":
    process_and_format()
